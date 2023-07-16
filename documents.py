import os
from datetime import date

from docx import Document

import openpyxl
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Font, Border, Side


class Reports:
    def __init__(self):
        self.__pathProcessing = "resources/documents/reportTemplate.docx"
        self.__pathNew = "resources/documents/reportTemplateNew.docx"
        self.__pathInWork = "resources/documents/reportTemplateInWork.docx"

    def create_user_app_report(self, number_app, fio, date_of_creation, dispatcher, executor, customer, address, description,
                               director_name, status):

        if status == 'Новая':
            report_docx = Document(self.__pathNew)
        elif status == 'В работе':
            report_docx = Document(self.__pathInWork)
        else:
            report_docx = Document(self.__pathProcessing)

        year, month, day = date_of_creation.split('-')

        data_dict = {"[number]": number_app, "[day]": day, "[month]": month, "[year]": year,
                     "[dispatcher]": dispatcher, "[executor]": executor, "[customer]": customer, "[fio]": fio,
                     "[address]": address, "[description]": description, "[director_name]": director_name}

        for paragraph in report_docx.paragraphs:
            for k, v in data_dict.items():
                paragraph.text = paragraph.text.replace(k, v)

        report_docx.save(f"resources/documents/report{number_app}.docx")

    def create_dispatcher_report(self):
        pass

    def create_work_for_apps_report(self, data):
        workbook = openpyxl.load_workbook("resources/documents/work_report.xlsx")
        worksheet = workbook.active

        style = Alignment(horizontal='center', vertical='center')
        font = Font(name='Arial', size=11)
        medium_border = Border(left=Side(style='medium'), right=Side(style='medium'),
                             top=Side(style='medium'), bottom=Side(style='medium'))

        lider_name = ''
        max_apps = 0

        for i, row in enumerate(data, 3):
            user_id, name, role, count_of_apps = row
            worksheet[f'A{i}'] = user_id
            worksheet[f'A{i}'].alignment = style
            worksheet[f'A{i}'].font = font
            worksheet[f'A{i}'].border = medium_border

            worksheet[f'B{i}'] = name
            worksheet[f'B{i}'].alignment = style
            worksheet[f'B{i}'].font = font
            worksheet[f'B{i}'].border = medium_border

            worksheet[f'C{i}'] = 'диспетчер' if role == 'dispatcher' else 'оператор'
            worksheet[f'C{i}'].alignment = style
            worksheet[f'C{i}'].font = font
            worksheet[f'C{i}'].border = medium_border

            worksheet[f'D{i}'] = count_of_apps
            worksheet[f'D{i}'].alignment = style
            worksheet[f'D{i}'].font = font
            worksheet[f'D{i}'].border = medium_border

            if count_of_apps > max_apps:
                max_apps = count_of_apps
                lider_name = name

        worksheet[f'F{3}'] = lider_name
        worksheet[f'F{3}'].alignment = style
        worksheet[f'F{3}'].font = font

        worksheet[f'G{3}'] = max_apps
        worksheet[f'G{3}'].alignment = style
        worksheet[f'G{3}'].font = font

        chart = BarChart()
        chart.type = "bar"
        chart.style = 11
        chart.title = "Гистограмма выполненных/обработанных заявок"
        chart.y_axis.title = 'Количество заявок'
        chart.x_axis.title = 'Данные менеджера'

        data = Reference(worksheet, min_col=4, min_row=3, max_row=len(data)+3)
        cats = Reference(worksheet, min_col=1, min_row=3, max_col=2, max_row=len(data)+3)

        chart.add_data(data)
        chart.set_categories(cats)
        chart.legend = None
        chart.shape = 4
        worksheet.add_chart(chart, "F6")

        workbook.save(
            f'resources/documents/work_report{str(date.today().day).zfill(2)}{str(date.today().month).zfill(2)}{str(date.today().year).zfill(2)}.xlsx')

    @staticmethod
    def open_file(file_name):
        full_path = os.path.abspath(file_name)
        os.startfile(full_path)
